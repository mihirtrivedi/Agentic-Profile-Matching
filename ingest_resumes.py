import os
import uuid
import docx
from langchain_core.documents import Document
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings

# 1. Generate 10 realistic sample resumes
resumes_data = [
    {"name": "Alice_Smith.txt", "content": "Alice Smith\nSoftware Engineer\nExperience: 5 years\nSkills: React, Node.js, AWS, TypeScript.\nDetails: Led frontend team to build scalable React applications. Deployed microservices on AWS using Node.js."},
    {"name": "Bob_Jones.docx", "content": "Bob Jones\nBackend Developer\nExperience: 3 years\nSkills: Python, Django, PostgreSQL, Docker.\nDetails: Built robust APIs using Django. Managed database schemas and optimized queries in PostgreSQL."},
    {"name": "Charlie_Brown.txt", "content": "Charlie Brown\nFull Stack Engineer\nExperience: 2 years\nSkills: React, Python, Flask, AWS.\nDetails: Developed user interfaces in React. Created backend services in Flask. Managed deployments on AWS EC2."},
    {"name": "Diana_Prince.docx", "content": "Diana Prince\nDevOps Engineer\nExperience: 6 years\nSkills: AWS, Kubernetes, Terraform, CI/CD.\nDetails: Architected cloud infrastructure on AWS. Automated deployments using Jenkins and Kubernetes."},
    {"name": "Evan_Wright.txt", "content": "Evan Wright\nFrontend Developer\nExperience: 4 years\nSkills: Vue.js, JavaScript, HTML, CSS.\nDetails: Designed responsive web pages. Worked closely with design teams to implement pixel-perfect UIs."},
    {"name": "Fiona_Gallagher.docx", "content": "Fiona Gallagher\nData Scientist\nExperience: 3 years\nSkills: Python, Pandas, Scikit-Learn, SQL.\nDetails: Developed machine learning models to predict customer churn. Processed large datasets using Pandas."},
    {"name": "George_Miller.txt", "content": "George Miller\nSoftware Engineer\nExperience: 7 years\nSkills: Java, Spring Boot, Microservices, AWS.\nDetails: Migrated monolithic applications to microservices using Spring Boot. Deployed on AWS ECS."},
    {"name": "Hannah_Abbott.docx", "content": "Hannah Abbott\nCloud Architect\nExperience: 8 years\nSkills: AWS, Azure, System Design, Node.js.\nDetails: Designed multi-cloud architectures. Built serverless applications using AWS Lambda and Node.js."},
    {"name": "Ian_Malcolm.txt", "content": "Ian Malcolm\nFull Stack Developer\nExperience: 5 years\nSkills: React, Node.js, MongoDB, Express.\nDetails: Built MERN stack applications from scratch. Implemented real-time features using WebSockets."},
    {"name": "Julia_Roberts.docx", "content": "Julia Roberts\nMobile Developer\nExperience: 4 years\nSkills: React Native, iOS, Android, Firebase.\nDetails: Developed cross-platform mobile apps using React Native. Integrated Firebase for real-time databases."},
    {"name": "Kevin_Hart.txt", "content": "Kevin Hart\nData Analyst\nExperience: 3 years\nSkills: SQL, Tableau, Excel, Python.\nDetails: Built complex dashboards in Tableau to track KPIs. Performed ad-hoc data analysis using SQL and Python."},
    {"name": "Laura_Croft.docx", "content": "Laura Croft\nCyber Security Analyst\nExperience: 5 years\nSkills: Network Security, Penetration Testing, Python, Kali Linux.\nDetails: Conducted vulnerability assessments and penetration testing. Implemented security protocols to protect internal networks."},
    {"name": "Michael_Scott.txt", "content": "Michael Scott\nProduct Manager\nExperience: 6 years\nSkills: Agile, Scrum, Jira, Product Strategy.\nDetails: Led cross-functional teams to deliver 3 major product launches. Defined product roadmaps and prioritized backlog."},
    {"name": "Nina_Simone.docx", "content": "Nina Simone\nUI/UX Designer\nExperience: 4 years\nSkills: Figma, Adobe XD, User Research, Wireframing.\nDetails: Designed user-centric interfaces for web and mobile. Conducted A/B testing and user research to improve conversion rates."},
    {"name": "Oscar_Martinez.txt", "content": "Oscar Martinez\nQA Automation Engineer\nExperience: 5 years\nSkills: Selenium, Cypress, Python, Test Automation.\nDetails: Built end-to-end testing frameworks from scratch using Cypress. Reduced manual testing time by 70%."},
    {"name": "Penny_Lane.docx", "content": "Penny Lane\nMachine Learning Engineer\nExperience: 4 years\nSkills: PyTorch, TensorFlow, Python, NLP.\nDetails: Trained large language models for sentiment analysis. Deployed ML models into production environments."},
    {"name": "Quentin_T.txt", "content": "Quentin T.\nC++ Developer\nExperience: 7 years\nSkills: C++, Qt, Embedded Systems, Linux.\nDetails: Developed high-performance embedded systems in C++. Optimized memory management for IoT devices."},
    {"name": "Rachel_Green.docx", "content": "Rachel Green\nFrontend Developer\nExperience: 3 years\nSkills: Angular, TypeScript, RxJS, HTML/CSS.\nDetails: Built single-page applications using Angular. Managed complex state using RxJS observables."},
    {"name": "Sam_Winchester.txt", "content": "Sam Winchester\nSite Reliability Engineer\nExperience: 5 years\nSkills: Linux, Bash, Prometheus, Grafana, Docker.\nDetails: Set up monitoring and alerting using Prometheus and Grafana. Managed incident response and improved system uptime."},
    {"name": "Tina_Fey.docx", "content": "Tina Fey\nBackend Developer\nExperience: 6 years\nSkills: Go, gRPC, PostgreSQL, Microservices.\nDetails: Architected high-throughput microservices in Go. Designed efficient communication protocols using gRPC."}
]

os.makedirs("sample_resumes", exist_ok=True)

for res in resumes_data:
    filepath = os.path.join("sample_resumes", res["name"])
    if res["name"].endswith(".txt"):
        with open(filepath, "w") as f:
            f.write(res["content"])
    elif res["name"].endswith(".docx"):
        doc = docx.Document()
        doc.add_paragraph(res["content"])
        doc.save(filepath)

print(f"Successfully generated {len(resumes_data)} sample resumes in 'sample_resumes/' directory.")

# 2. Ingest into Chroma Database
print("Initializing HuggingFace Embeddings...")
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

print("Connecting to local Chroma Database...")
vectorstore = Chroma(persist_directory="./chroma_db", embedding_function=embeddings)

documents = []
for res in resumes_data:
    doc = Document(
        page_content=res["content"],
        metadata={"source": res["name"], "candidate_id": str(uuid.uuid4())[:8]}
    )
    documents.append(doc)

print("Ingesting resumes into Chroma vector database...")
vectorstore.add_documents(documents)

print("SUCCESS! The database is now populated and ready for your corporate demo.")
